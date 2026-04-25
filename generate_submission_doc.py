"""Generate the SHAKTI AI EliteHer hackathon submission as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PRIMARY = RGBColor(0x63, 0x0E, 0xD4)   # deep purple
ACCENT = RGBColor(0xB4, 0x13, 0x6D)    # vibrant pink
DARK = RGBColor(0x18, 0x14, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x66)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_mono(doc, text, size=9):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 4"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(hdr[i], "630ED4")

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4136D")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ============================================================
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============== TITLE PAGE ==============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SHAKTI AI")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = PRIMARY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("The Digital Sanctuary")
run.italic = True
run.font.size = Pt(18)
run.font.color.rgb = ACCENT

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run(
    "A zero-cost, AI-powered safety, wellness, and empowerment platform for women in India."
)
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = MUTED

doc.add_paragraph()
event = doc.add_paragraph()
event.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = event.add_run("EliteHer Hackathon Submission")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = DARK

add_divider(doc)
doc.add_paragraph()

# ============== 1. IDEA TITLE ==============
add_heading(doc, "1. Idea Title", level=1, color=PRIMARY)
add_para(
    doc,
    "SHAKTI AI — The Digital Sanctuary",
    bold=True,
    size=14,
    color=ACCENT,
)
add_para(
    doc,
    "A zero-cost, AI-powered safety, wellness, and empowerment platform for women in India.",
    italic=True,
    color=MUTED,
)

# ============== 2. IDEA DESCRIPTION ==============
add_heading(doc, "2. Idea Description", level=1, color=PRIMARY)
add_para(
    doc,
    "In India, women face fragmented digital tools — one app for safety, another for periods, "
    "another for jobs, another for mental health. None of them talk to each other. None of them are "
    "designed for her holistically.",
)
p = doc.add_paragraph()
run = p.add_run("SHAKTI AI")
run.bold = True
run.font.color.rgb = ACCENT
p.add_run(
    " unifies six critical pillars of a woman's life — Safety, Health, Careers, Tech Upskilling, "
    "Community, and Identity — into one cohesive 'Digital Sanctuary.' Every feature is powered by "
    "free-tier AI (no paid APIs, no subscription walls), making it accessible to women across every "
    "economic strata."
)

p = doc.add_paragraph()
run = p.add_run("The impact: ")
run.bold = True
p.add_run(
    "A woman in a Tier-3 town can scan a packaged snack to check if it disrupts her hormones, get an "
    "AI-generated 6-month career plan to land a remote job, talk to a compassionate mental-health "
    "companion at 3 AM, and trigger an SOS that broadcasts her live location — all in the same app, "
    "all for free."
)

# ============== 3. TECHNICAL DETAILS ==============
add_heading(doc, "3. Technical Details", level=1, color=PRIMARY)

# 3a. Technologies Used
add_heading(doc, "3a. Technologies Used", level=2, color=ACCENT)
add_table(
    doc,
    headers=["Layer", "Technology"],
    rows=[
        ["Frontend Framework", "React 19 + Vite 8 (SPA)"],
        ["Routing", "React Router 7"],
        ["State Management", "Zustand (5 isolated stores: auth, health, jobs, safety, ui)"],
        ["Styling", "Tailwind CSS 4 + custom glassmorphism CSS variables"],
        ["Animation", "Framer Motion 12"],
        ["Maps", "React-Leaflet + OpenStreetMap (CartoDB Voyager tiles)"],
        ["Charts", "Recharts 3 (mood / cycle trend graphs)"],
        ["OCR", "Tesseract.js (offline label reading fallback)"],
        ["Notifications", "react-hot-toast"],
        ["Icons", "Lucide React"],
    ],
)
doc.add_paragraph()

# 3b. Architecture Overview
add_heading(doc, "3b. Architecture Overview", level=2, color=ACCENT)
arch = """┌──────────────────────────────────────────────────┐
│           CLIENT (React SPA on Netlify)          │
│                                                  │
│  ┌────────────┐  ┌────────────┐  ┌────────────┐  │
│  │   Pages    │  │ Components │  │  Zustand   │  │
│  │  (6 hubs)  │◄─┤  (shared)  │◄─┤  Stores    │  │
│  └─────┬──────┘  └────────────┘  └─────┬──────┘  │
│        │                                │        │
│        ▼                                ▼        │
│  ┌──────────────┐              ┌────────────────┐│
│  │ aiService.js │              │  firebase.js   ││
│  │ (pure fetch) │              │ (Auth + Store) ││
│  └──────┬───────┘              └────────┬───────┘│
└─────────┼───────────────────────────────┼────────┘
          │                               │
          ▼                               ▼
   ┌─────────────┐               ┌──────────────┐
   │ OpenRouter  │               │   Firebase   │
   │   (8+ AI    │               │   Firestore  │
   │   models)   │               │   + Auth     │
   └─────────────┘               └──────────────┘
          │
          ▼
   ┌─────────────────────────────────┐
   │  Cascading Fallback Strategy:   │
   │  Vision: Gemma-12B → 27B → 4B   │
   │  Text:   MiniMax → Gemma → Llama│
   └─────────────────────────────────┘"""
add_mono(doc, arch, size=8)

add_para(doc, "Key architectural decisions:", bold=True)
add_bullets(
    doc,
    [
        "Pure fetch AI layer — no SDK lock-in; we can swap any model at any time.",
        "Multi-model fallback chains — if one free model rate-limits, we cascade to the next, guaranteeing uptime.",
        "Client-side image compression (max 800px @ 70% JPEG quality) before API upload — reduces payload by ~80%, speeds up vision calls.",
        "Optimistic UI — Zustand updates state instantly; Firebase syncs in the background.",
    ],
)

# 3c. Database
add_heading(doc, "3c. Database Used", level=2, color=ACCENT)
p = doc.add_paragraph()
run = p.add_run("Firebase Firestore")
run.bold = True
p.add_run(" (NoSQL document store, free Spark tier) with custom security rules in ")
run = p.add_run("firestore.rules")
run.font.name = "Consolas"
p.add_run(".")

add_para(doc, "Collections:", bold=True)
add_bullets(
    doc,
    [
        "users — profile, ShePoints, onboarding state",
        "moodLogs — daily emotion entries",
        "cycleData — menstrual tracking",
        "incidents — Evidence Locker entries",
        "forums, events — community content",
        "trustedCircle — emergency contact relationships",
    ],
)
add_para(
    doc,
    "Firebase Auth handles email/password authentication, with protected routes via ProtectedRoute.jsx.",
)

# 3d. Third-Party Integrations
add_heading(doc, "3d. Third-Party Integrations — The API Spine", level=2, color=ACCENT)
add_para(
    doc,
    "Every 'smart' feature in SHAKTI is powered by a deliberate API choreography. "
    "Total cost of all integrations: ₹0/month.",
    italic=True,
    color=MUTED,
)

# Integration 1
add_heading(doc, "Integration #1 — OpenRouter API (The AI Backbone)", level=3)
add_para(
    doc,
    "OpenRouter aggregates 50+ free open-source LLMs behind a single OpenAI-compatible endpoint. "
    "We use it as our entire AI brain.",
)
add_bullets(
    doc,
    [
        "Endpoint: https://openrouter.ai/api/v1/chat/completions",
        "Auth: Bearer token in env var VITE_OPENROUTER_API_KEY",
        "Implementation: src/services/aiService.js",
    ],
)
add_para(doc, "9 distinct AI features wired through OpenRouter:", bold=True)
add_table(
    doc,
    headers=["#", "Feature", "Model", "What it does"],
    rows=[
        ["1", "AI Companion (Mental Health Chatbot)", "minimax/minimax-m2.5:free",
         "Compassionate conversation; auto-injects crisis helplines (iCall, AASRA, Women Helpline 181) when self-harm keywords are detected."],
        ["2", "Skill-to-Income Translator", "minimax/minimax-m2.5:free",
         "Takes user skills + location → returns 5–6 income opportunities in INR with platforms (Vedantu, Upwork, Swiggy Home Chef)."],
        ["3", "Career Growth Simulator", "minimax/minimax-m2.5:free",
         "Generates a structured 6-month JSON plan with monthly milestones, salary progression, and free course recommendations."],
        ["4", "AI Project Generator", "minimax/minimax-m2.5:free",
         "Suggests buildable tech projects (1–4 weeks) with social impact and tech stack."],
        ["5", "Emotion Analysis", "minimax/minimax-m2.5:free",
         "Sentiment + intensity scoring with needsSupport boolean for escalation."],
        ["6", "Safety Route Analyzer", "minimax/minimax-m2.5:free",
         "Origin + destination + time-of-day → safety score (1–10), risk factors, practical tips."],
        ["7", "Wellness Activity Engine", "minimax/minimax-m2.5:free",
         "Mood + energy + time-available → tailored breathing/stretch/journaling exercise."],
        ["8", "Mentor Match", "minimax/minimax-m2.5:free",
         "Profile → ideal mentor archetype + suggested topics + meeting cadence."],
        ["9", "'Is This Normal?' Incident Analyzer", "minimax/minimax-m2.5:free",
         "Legal/psychological triage of harassment incidents → category, severity, action steps, helplines."],
    ],
)
doc.add_paragraph()

# Integration 2 - Food Scanner
add_heading(doc, "Integration #2 — Multi-Stage Vision Pipeline (Hormone-Safe Food Scanner)", level=3)
add_para(
    doc,
    "Our most technically sophisticated integration. The Food Scanner runs a 3-step orchestration "
    "across two different model families.",
)
add_para(doc, "Implementation: src/pages/health/FoodScanner.jsx", italic=True, color=MUTED)

pipeline = """User uploads photo
       │
       ▼
[1] Client-side image compression (800px, 70% JPEG)
       │
       ▼
[2] VISION CALL → OpenRouter
    Model cascade: gemma-3-12b-it → gemma-3-27b-it → gemma-3-4b-it
    Prompt: "List every food item, drink, snack, or ingredient.
             If it's a packaged label, extract all ingredients."
    Output: comma-separated food list
       │
       ▼
[3] ANALYSIS CALL → OpenRouter
    Model cascade: minimax-m2.5 → gemma-3-4b → llama-4-scout
    Prompt: "Analyze for women's hormonal health — PCOS, estrogen
             dominance, thyroid, endocrine disruption.
             Return JSON {ingredients, flagged, safeAlternatives}"
    Output: structured JSON
       │
       ▼
UI renders flagged items with severity (High/Medium/Low)
+ awards 15 ShePoints to user"""
add_mono(doc, pipeline, size=8)

add_para(doc, "Why this matters technically:", bold=True)
add_bullets(
    doc,
    [
        "Two completely different model architectures working in series (vision → text reasoning).",
        "Resilience via cascading fallbacks — 6 models total; the app never breaks if one rate-limits.",
        "30-second timeout with AbortController prevents hung requests.",
        "JSON repair logic with regex fallbacks for malformed model output.",
    ],
)

# Integration 3
add_heading(doc, "Integration #3 — Firebase Suite", level=3)
add_table(
    doc,
    headers=["Service", "Use"],
    rows=[
        ["Firebase Auth", "Email/password signup, session persistence, protected routes."],
        ["Firestore", "Real-time NoSQL database (mood logs, cycle data, evidence locker, community posts)."],
        ["Firestore Security Rules", "Per-user data isolation — defined in firestore.rules."],
        ["Firebase Hosting", "Optional secondary deploy target."],
    ],
)
doc.add_paragraph()

# Integration 4
add_heading(doc, "Integration #4 — Mapping Stack (OpenStreetMap)", level=3)
add_para(
    doc,
    "No Google Maps, no Mapbox keys. We use the fully free OSM ecosystem:",
)
add_bullets(
    doc,
    [
        "React-Leaflet for the React component layer.",
        "CartoDB Voyager tiles for clean styled basemaps.",
        "OpenStreetMap data for points of interest (hospitals, police stations).",
        "Used in: LiveTracking.jsx, SafetyMap.jsx",
    ],
)

# Integration 5
add_heading(doc, "Integration #5 — Tesseract.js (On-Device OCR Fallback)", level=3)
add_para(
    doc,
    "Pure-JavaScript OCR running in the browser via WebAssembly. Used as an offline fallback for the "
    "Food Scanner when the user is on a poor connection — extracts label text locally with no API call.",
)

# Integration 6
add_heading(doc, "Integration #6 — Netlify (Hosting + CI/CD)", level=3)
add_bullets(
    doc,
    [
        "Automated deploys via netlify deploy CLI from dist/ build output.",
        "public/_redirects handles SPA fallback routing for React Router.",
        "Project ID: b8f87d52-570a-497c-9ba9-e73de837389a",
        "Edge functions ready for future server-side needs.",
    ],
)

# ============== 4. SUBMISSION LINKS ==============
add_heading(doc, "4. Submission Links", level=1, color=PRIMARY)
add_table(
    doc,
    headers=["Link", "URL"],
    rows=[
        ["GitHub Repository (Primary)", "https://github.com/DuvvuruDeepakReddy18/Shakti-Ai"],
        ["GitHub Repository (Mirror)", "https://github.com/BANU-09/Shakti-Ai"],
        ["Live Demo (Netlify)", "[Add your Netlify URL — e.g., https://shakti-ai.netlify.app]"],
        ["Figma / Design Mockups", "[Add link — stitch-designs/ folder has source PNGs]"],
        ["Pitch Deck (PPT)", "[Add link]"],
    ],
)
doc.add_paragraph()
add_para(
    doc,
    "Action item: Set all repos and links to public access. Verify the .env is gitignored "
    "(it is) so no secrets leak.",
    italic=True,
    color=ACCENT,
)

# ============== 5. DEMO VIDEO SCRIPT ==============
add_heading(doc, "5. Demo Video Script (3–5 minutes)", level=1, color=PRIMARY)

add_heading(doc, "Opening (0:00 – 0:20)", level=3)
add_para(
    doc,
    "\"Meet Priya. She's 24, lives in Hyderabad, just landed her first job. In the next 5 minutes, "
    "watch how SHAKTI AI becomes her co-pilot for safety, health, and career — without costing her "
    "or us a single rupee in API fees.\"",
    italic=True,
)

add_heading(doc, "Product Walkthrough (0:20 – 3:30)", level=3)

add_para(doc, "Scene 1 — Onboarding & Landing (20s)", bold=True)
add_bullets(doc, ["Show login → onboarding wizard → dashboard with 6 module cards."])

add_para(doc, "Scene 2 — Hormone-Safe Food Scanner (45s) ★ HERO FEATURE", bold=True, color=ACCENT)
add_bullets(
    doc,
    [
        "Snap a photo of a packet of chips.",
        "Watch the 3-step pipeline: compression → Gemma vision identifies 'potato chips, MSG, palm oil' → MiniMax flags 'MSG: High risk for thyroid; Palm oil: Medium risk for inflammation.'",
        "Show safe alternatives + ShePoints reward.",
    ],
)

add_para(doc, "Scene 3 — AI Companion (40s)", bold=True)
add_bullets(
    doc,
    [
        "Type 'I've been feeling overwhelmed at work.'",
        "Show real-time empathetic response with breathing exercise.",
        "Mention crisis-keyword auto-escalation to helplines.",
    ],
)

add_para(doc, "Scene 4 — SOS + Safety Map (40s)", bold=True)
add_bullets(
    doc,
    [
        "Tap the global pulsing SOS button.",
        "Show emergency overlay → live tracking on OpenStreetMap.",
        "Display nearest safe zones, hospitals.",
    ],
)

add_para(doc, "Scene 5 — Career Simulator (35s)", bold=True)
add_bullets(
    doc,
    [
        "Input: 'HTML, JavaScript' → target 'Frontend Developer'.",
        "AI generates a 6-month JSON plan with salary curve from ₹15K → ₹45K.",
        "Render plan on Recharts visualization.",
    ],
)

add_para(doc, "Scene 6 — Skill Translator (30s)", bold=True)
add_bullets(
    doc,
    [
        "Input 'cooking, teaching, English'.",
        "Get 5 income opportunities with INR estimates and platforms.",
    ],
)

add_heading(doc, "Technical Highlights (3:30 – 4:30)", level=3)
add_bullets(
    doc,
    [
        "Show aiService.js — 'Pure fetch, 9 features, one file, zero SDK dependencies.'",
        "Show fallback chain logs in console — 'Watch how we cascade through 3 vision models when one rate-limits.'",
        "Mention: 100% free-tier infrastructure, deployable for $0/month.",
    ],
)

add_heading(doc, "Real-Use Scenario (4:30 – 5:00)", level=3)
add_para(
    doc,
    "\"A college student in Vizag uses the food scanner before buying groceries; an HR exec in "
    "Bangalore uses 'Is This Normal?' to triage a harassment complaint; a single mother in Pune "
    "uses the Skill Translator to find tutoring gigs. One app. One sanctuary. Every woman.\"",
    italic=True,
)

# ============== 6. PITCH VIDEO SCRIPT ==============
add_heading(doc, "6. Pitch Video Script (2–3 minutes)", level=1, color=PRIMARY)

add_heading(doc, "Problem (0:00 – 0:30)", level=3)
add_para(
    doc,
    "\"1 in 3 women in India experience harassment. 80% experience hormone-related health issues "
    "like PCOS. 70% of working women want to upskill but can't afford courses. They open 7 different "
    "apps to manage these — none of which understand them as whole humans.\"",
    italic=True,
)

add_heading(doc, "Solution (0:30 – 1:15)", level=3)
add_para(
    doc,
    "\"SHAKTI AI is one platform, six ecosystems — Safety, Health, Careers, Tech, Community, "
    "Identity — bound together by a custom AI layer. Every feature, from our hormone-safe food "
    "scanner to our 24/7 mental-health companion, runs on free-tier APIs. We engineered cascading "
    "model fallbacks so the app never breaks, even at scale. Built on React, Firebase, OpenRouter, "
    "OpenStreetMap, and Netlify — total infra cost: zero.\"",
    italic=True,
)

add_heading(doc, "Market & Impact (1:15 – 2:00)", level=3)
add_bullets(
    doc,
    [
        "TAM: 680 million women in India, 350M+ smartphone users.",
        "SAM: 120M urban + semi-urban women aged 18–45 with internet access.",
        "Beachhead: 5M women in Tier-2 / Tier-3 cities seeking safety + income tools.",
        "Impact metric: Each user touched = 1 less app to juggle, 1 less paywall, 1 more decision made with AI assistance.",
        "Path to revenue: Freemium ShePoints economy → premium therapist matching, certified courses, brand-safe sponsored skill gigs.",
    ],
)

add_heading(doc, "Why Our Team (2:00 – 2:45)", level=3)
add_para(
    doc,
    "\"[Names of team members] — we built this in [X weeks] because [personal motivation]. We're "
    "not just engineers — we're the daughters, sisters, and friends of the women this serves. "
    "We've shipped a working product, not a slide deck. Try it, break it, give us feedback. We'll "
    "keep building.\"",
    italic=True,
)

add_divider(doc)
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("SHAKTI AI · The Digital Sanctuary · Built for EliteHer Hackathon")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = MUTED

# Save
output_path = "SHAKTI_AI_Submission.docx"
doc.save(output_path)
print(f"Generated: {output_path}")
